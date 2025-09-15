import html
import re
from docx import Document
from bs4 import BeautifulSoup as BS
from lxml import etree
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Pt, Inches, Cm, RGBColor
import pandas as pd
import numpy as np
from latex2mathml import converter as latex_converter
from io import BytesIO

PAGE_SIZE = {
    "portrait": {
        "orientation": WD_ORIENT.PORTRAIT,
        "x": 21,
        "y": 29.7
    },
    "landscape": {
        "orientation": WD_ORIENT.LANDSCAPE,
        "x": 29.7,
        "y": 21
    }
}

hAlignments = {
    0: 'left',
    1: 'center',
    2: 'right',
    3: 'justify',
}

vAlignments = {
    0: 'top',
    1: 'center',
    3: 'bottom',
}

namespaces = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
}

horizontal_alignments = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
table_alignments = {
    "left": WD_TABLE_ALIGNMENT.LEFT,
    "right": WD_TABLE_ALIGNMENT.RIGHT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
}

vertical_alignments = {
    "top": WD_ALIGN_VERTICAL.TOP,
    "center": WD_ALIGN_VERTICAL.CENTER,
    "bottom": WD_ALIGN_VERTICAL.BOTTOM,
}


class Json2Docx:
    def process_block(self, data):
        optized_data = []
        grouped_data = []
        new_group = []
        prev = {}
        labels = []
        positions = []
        alignments = []
        types = []
        footers = []
        for c, item in enumerate(data):
            id = prev.get("ID", None)
            label = prev.get("label", None)
            if (not prev or id == item["ID"]) and not label == "footer":
                new_group.append(item)
                labels.append(item["label"])
                positions.append(item["position"])
                alignments.append(item["alignment"])
                types.append(item["type"])

            else:
                grouped_data.append(
                    {"contents": new_group, "labels": labels, "positions": positions, "alignments": alignments,
                     "types": types})
                new_group = [item]
                labels = [item["label"]]
                positions = [item["position"]]
                alignments = [item["alignment"]]
                types = [item["type"]]

            if c == len(data) - 1:
                grouped_data.append(
                    {"contents": new_group, "labels": labels, "positions": positions, "alignments": alignments,
                     "types": types})

            prev = item

        for groups in grouped_data:
            prevGrID = None
            page_size = groups["contents"][0]["page_size"]
            page = {"number": groups["contents"][0]["page"], "width": page_size["width"], "height": page_size["height"],
                    "ratio": {
                        "width": (PAGE_SIZE["portrait"]["x"] * 96 / 2.54) / page_size["width"],
                        "height": (PAGE_SIZE["portrait"]["y"] * 96 / 2.54) / page_size["height"]
                    }, "orientation": "portrait"}

            if page["width"] > page["height"]:
                page["ratio"] = {
                    "width": (PAGE_SIZE["portrait"]["x"] * 96 / 2.54) / page_size["width"],
                    "height": (PAGE_SIZE["portrait"]["y"] * 96 / 2.54) / page_size["height"]
                }
                page["orientation"] = "landscape"

            block = {"page": page, "block_labels": groups["labels"], "block_positions": groups["positions"],
                     "block_alignments": groups["alignments"], "block_types": groups["types"]}

            block["groups"] = []
            new_group = []
            if groups["labels"][0] == "footer":
                block["groups"].append(groups["contents"])
                footers.append(block)
                continue

            for c, item in enumerate(groups["contents"]):

                if prevGrID is None or item["grID"] == prevGrID:
                    new_group.append(item)

                else:
                    block["groups"].append(new_group)
                    new_group = [item]

                if c == len(groups["contents"]) - 1:
                    block["groups"].append(new_group)

                prevGrID = item["grID"]

            optized_data.append(block)

        return optized_data, footers

    def single_group(self, data, prevData, nextData):
        groups = data["groups"]
        label = data["block_labels"][0]
        page = data["page"]
        item = groups[0][0]
        position = item["position"] if item["position"] is not None else "left"

        x0, y0, x1, y1 = item["box"]["xmin"], item["box"]["ymin"], item["box"]["xmax"], item["box"]["ymax"]

        if item["type"] == "image":
            width = (x1 - x0) * page["ratio"]["width"] * 0.75
            height = (y1 - y0) * page["ratio"]["height"] * 0.75
            p = self.document.add_paragraph()
            p_format = p.paragraph_format
            run = p.add_run()
            inline_shape = run.add_picture(item["src"])
            inline_shape.width = Pt(width)
            inline_shape.height = Pt(height)
            p_format.alignment = horizontal_alignments[item["alignment"]]

        elif label == "paragraph":
            p = self.document.add_paragraph()
            p_format = p.paragraph_format
            run = p.add_run()
            font = run.font

            run.text = html.unescape(item['src'])
            p_format.first_line_indent = Inches(0.5)
            p_format.alignment = horizontal_alignments[item["alignment"]]
            font.bold = item["font"]["weight"] == "bold"
            font.italic = item["font"]["style"] == "italic"
            font.size = Pt(item["font"]["size"])

        elif label == "continued" or label == "lined":
            p = self.document.add_paragraph()
            run = p.add_run()
            font = run.font

            run.text = html.unescape(item['src'])
            font.size = Pt(item["font"]["size"])


        elif label == "caption":
            cap_table = self.document.add_table(rows=1, cols=1)
            cap_table.alignment = table_alignments[position]

            cell = cap_table.cell(0, 0)
            p = cell.paragraphs[0]
            p_format = p.paragraph_format
            run = p.add_run(html.unescape(item['src']))
            font = run.font

            run.text = item["src"]
            font.bold = item["font"]["weight"] == "bold"
            font.italic = item["font"]["style"] == "italic"
            font.size = Pt(item["font"]["size"])
            cap_table.cell(0, 0).vertical_alignment = vertical_alignments["center"]

            if nextData["block_labels"][0] == "table" or nextData["block_labels"][0] == "picture":
                width = nextData["groups"][0][0]["box"]["xmax"] - nextData["groups"][0][0]["box"]["xmin"]
                ratio = nextData["page"]["ratio"]
                cap_table.cell(0, 0).width = Pt(width * ratio["width"] * 0.75)
                p_format.space_after = 0
                p_format.alignment = horizontal_alignments[position]
            else:
                # cap_table.cell(0, 0).width = Pt((x1 - x0) * page["ratio"]["width"] * 0.75)
                cap_table.cell(0, 0).width = Inches(3)
                p_format.alignment = horizontal_alignments[item["alignment"]]

        elif label == "subtitle":
            p = self.document.add_paragraph()
            p_format = p.paragraph_format
            run = p.add_run()
            font = run.font

            run.text = html.unescape(item['src'])
            p_format.first_line_indent = Inches(0.5)
            p_format.alignment = horizontal_alignments[item["alignment"]]
            font.bold = item["font"]["weight"] == "bold"
            font.italic = item["font"]["style"] == "italic"
            font.size = Pt(item["font"]["size"])

        elif label == "title" or label == "central":
            p = self.document.add_paragraph()
            p_format = p.paragraph_format
            run = p.add_run()
            font = run.font

            run.text = html.unescape(item['src'])
            p_format.alignment = horizontal_alignments[item["alignment"]]
            font.bold = item["font"]["weight"] == "bold"
            font.italic = item["font"]["style"] == "italic"
            font.size = Pt(item["font"]["size"])

        elif label == "formula":
            try:
                pattern = r"\\eqno|\\eqn|\\left|\\right|\\Bigg \\|\\left\\|\\right\\"
                updated_latex_string = re.sub(pattern, " ", item["src"]["latex"])
                mathml_code = latex_converter.convert(updated_latex_string)
                tree = etree.fromstring(mathml_code)
                xslt = etree.parse("MML2OMML.XSL")
                transform = etree.XSLT(xslt)
                new_dom = transform(tree)

                p = self.document.add_paragraph()

                p_format = p.paragraph_format
                font = p.style.font
                p._element.append(new_dom.getroot())

                p_format.alignment = horizontal_alignments[item["alignment"]]
                p.bold = item["font"]["weight"] == "bold"
                font.italic = item["font"]["style"] == "italic"
                font.size = Pt(item["font"]["size"])

            except:
                width = (x1 - x0) * page["ratio"]["width"] * 0.75
                height = (y1 - y0) * page["ratio"]["height"] * 0.75
                p = self.document.add_paragraph()
                p_format = p.paragraph_format
                run = p.add_run()
                inline_shape = run.add_picture(item["src"]["image"])
                inline_shape.width = Pt(width)
                inline_shape.height = Pt(height)
                p_format.alignment = horizontal_alignments[item["alignment"]]


        elif label == "table":
            width = (x1 - x0) * page["ratio"]["width"] * 0.75
            self.process_table(item, self.document, table_width=width)
            self.document.add_paragraph()

        else:
            print("Label not Found", label)

    def multiple_group(self, data):
        groups = data["groups"]
        page = data["page"]
        table = self.document.add_table(1, len(groups))

        for g, group in enumerate(groups):
            cell = table.cell(0, g)
            cell.vertical_alignment = vertical_alignments["center"]
            for i, item in enumerate(group):
                x0, y0, x1, y1 = item["box"]["xmin"], item["box"]["ymin"], item["box"]["xmax"], item["box"]["ymax"]
                if item["type"] == "image":
                    width = (x1 - x0) * page["ratio"]["width"] * 0.75
                    height = (y1 - y0) * page["ratio"]["height"] * 0.75
                    p = cell.add_paragraph()
                    p_format = p.paragraph_format
                    run = p.add_run()
                    
                    inline_shape = run.add_picture(item["src"])
                    inline_shape.width = Pt(width)
                    inline_shape.height = Pt(height)
                    p_format.alignment = horizontal_alignments[item["alignment"]]

                elif item["label"] == "paragraph":
                    p = cell.add_paragraph()
                    p_format = p.paragraph_format
                    run = p.add_run()
                    font = run.font

                    run.text = html.unescape(item['src'])
                    p_format.alignment = horizontal_alignments[item["alignment"]]
                    font.bold = item["font"]["weight"] == "bold"
                    font.italic = item["font"]["style"] == "italic"
                    font.size = Pt(item["font"]["size"])

                elif item["label"] == "continued" or item["label"] == "lined":
                    p = cell.paragraphs[0]
                    run = p.add_run()
                    font = run.font

                    run.text = html.unescape(item['src'])
                    font.size = Pt(item["font"]["size"])

                elif item["label"] == "caption":
                    p = cell.paragraphs[0]
                    p_format = p.paragraph_format
                    run = p.add_run()
                    font = run.font

                    run.text = html.unescape(item['src'])
                    p_format.alignment = horizontal_alignments[item["alignment"]]
                    font.bold = item["font"]["weight"] == "bold"
                    font.italic = item["font"]["style"] == "italic"
                    font.size = Pt(item["font"]["size"])

                elif item["label"] == "subtitle":
                    p = cell.paragraphs[0]
                    p_format = p.paragraph_format
                    run = p.add_run()
                    font = run.font

                    run.text = html.unescape(item['src'])
                    p_format.alignment = horizontal_alignments[item["alignment"]]
                    font.bold = item["font"]["weight"] == "bold"
                    font.italic = item["font"]["style"] == "italic"
                    font.size = Pt(item["font"]["size"])

                elif item["label"] == "title" or item["label"] == "central":
                    p = cell.paragraphs[0]
                    p_format = p.paragraph_format
                    run = p.add_run()
                    font = run.font

                    run.text = html.unescape(item['src'])
                    p_format.alignment = horizontal_alignments[item["alignment"]]
                    font.bold = item["font"]["weight"] == "bold"
                    font.italic = item["font"]["style"] == "italic"
                    font.size = Pt(item["font"]["size"])

                elif item["label"] == "formula":
                    try:
                        pattern = r"\\eqno|\\eqn|\\left|\\right|\\Bigg \\|\\left\\|\\right\\"
                        updated_latex_string = re.sub(pattern, " ", item["src"]["latex"])
                        mathml_code = latex_converter.convert(updated_latex_string)
                        tree = etree.fromstring(mathml_code)
                        xslt = etree.parse("MML2OMML.XSL")
                        transform = etree.XSLT(xslt)
                        new_dom = transform(tree)
                        p = cell.paragraphs[0]
                        p_format = p.paragraph_format
                        font = p.style.font
                        p._element.append(new_dom.getroot())
                        p_format.alignment = horizontal_alignments[item["alignment"]]
                        font.bold = item["font"]["weight"] == "bold"
                        font.italic = item["font"]["style"] == "italic"
                        font.size = Pt(item["font"]["size"])
                    except:
                        width = (x1 - x0) * page["ratio"]["width"] * 0.75
                        height = (y1 - y0) * page["ratio"]["height"] * 0.75
                        p = cell.add_paragraph()
                        p_format = p.paragraph_format
                        run = p.add_run()
                        
                        inline_shape = run.add_picture(item["src"]["image"])
                        inline_shape.width = Pt(width)
                        inline_shape.height = Pt(height)
                        p_format.alignment = horizontal_alignments[item["alignment"]]


                elif item["label"] == "table":
                    width = (x1 - x0) * page["ratio"]["width"] * 0.75
                    self.process_table(item, self.document, table_width=width)

                else:
                    print("Label not Found", item["label"])


    def process_table(self, item, document, table_width):
        soup = BS(item["src"], 'html.parser')
        rows = soup.find_all('tr')
        cells = rows[0].find_all(['td'])
        max_cols = sum(int(cell.get('colspan', 1)) for cell in cells)
        max_rows = len(rows)
        matrix = [[np.empty for _ in range(max_cols)] for _ in range(max_rows)]
        spans = []

        for i, row in enumerate(rows):
            cells = row.find_all(['td'])

            for j, cell in enumerate(cells):
                text = cell.decode_contents()
                rowspan = int(cell.get('rowspan', 1))
                colspan = int(cell.get('colspan', 1))
                executed = False

                for r in range(max_rows):
                    for c in range(max_cols):
                        if matrix[r][c] == np.empty and matrix[r][c] != np.nan:
                            for ir in range(r, r + rowspan):
                                for jc in range(c, c + colspan):

                                    if ir == r and jc == c:
                                        matrix[ir][jc] = text

                                        spans.append({
                                            "base_index": {"x": jc, "y": ir},
                                            "merge_index": {"x": colspan + jc - 1, "y": rowspan + ir - 1},
                                        })

                                    else:
                                        matrix[ir][jc] = np.nan

                            executed = True

                        if executed:
                            break
                    if executed:
                        break

        df = pd.DataFrame(matrix)
        data = df.values.tolist()

        table = document.add_table(rows=max_rows, cols=max_cols)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        table.alignment = table_alignments[item["alignment"]]

        for i, row in enumerate(data):
            for j, cell in enumerate(row):
                table.cell(i, j).width = Pt(table_width) / max_cols
                if type(cell) == float and str(cell) == "nan":
                    continue
                text = cell.replace("<br/>", "\n")
                table_p = table.cell(i, j).paragraphs[0]
                table_p.paragraph_format.space_before = Pt(4)
                table_p.paragraph_format.space_after = Pt(4)
                table_p.paragraph_format.left_indent = Pt(4)
                table_p.paragraph_format.right_indent = Pt(4)
                run = table_p.add_run(text.strip())
                font = run.font
                font.bold = item["font"]["weight"] == "bold"
                font.italic = item["font"]["style"] == "italic"
                font.size = Pt(item["font"]["size"])

                table_p.alignment = horizontal_alignments["center"]
                table.cell(i, j).vertical_alignment = vertical_alignments["center"]

        for i, span in enumerate(spans):
            a = table.cell(span["base_index"]["y"], span["base_index"]["x"])
            b = table.cell(span["merge_index"]["y"], span["merge_index"]["x"])
            a.merge(b)

    def generate_docx(self, jsonContent):
        self.document = Document()
        self.style = self.document.styles["Normal"]
        self.style.font.name = "Times New Roman"
        self.style.font.size = Pt(14)

        for section in self.document.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)

        isCover = True
        orientation = None
        prevData = None
        nextData = None

        optimized_data, footers = self.process_block(jsonContent)

        for d, data in enumerate(optimized_data):
            page = data["page"]
            if page["number"] == 1 and "paragraph" in data["block_labels"]:
                isCover = False

            if page["number"] > 1 and isCover:
                self.document.add_page_break()
                isCover = False

            if d + 1 < len(optimized_data):
                nextData = optimized_data[d + 1]

            if orientation != page["orientation"]:
                if orientation is None:
                    section = self.document.sections[-1]
                    section.orientation = PAGE_SIZE[page["orientation"]]["orientation"]
                    section.page_width = Cm(PAGE_SIZE[page["orientation"]]["x"])
                    section.page_height = Cm(PAGE_SIZE[page["orientation"]]["y"])
                else:
                    section = self.document.add_section(WD_SECTION.NEW_PAGE)
                    section.orientation = PAGE_SIZE[page["orientation"]]["orientation"]
                    section.page_width = Cm(PAGE_SIZE[page["orientation"]]["x"])
                    section.page_height = Cm(PAGE_SIZE[page["orientation"]]["y"])

                orientation = page["orientation"]

            if len(data["block_labels"]) == 1:
                self.single_group(data, prevData, nextData)

            else:
                self.multiple_group(data)

            prevData = data

        if len(footers) > 0:
            footer_table = self.document.add_table(rows=len(footers), cols=1)

            for f, footer in enumerate(footers):
                groups = footer["groups"]
                item = groups[0][0]

                cell = footer_table.cell(f, 0)

                cell.text = ""

                p = cell.paragraphs[0]
                p_format = p.paragraph_format
                p_format.alignment = horizontal_alignments[item["alignment"]]

                run = p.add_run(html.unescape(item["src"]))
                font = run.font
                font.bold = item["font"]["weight"] == "bold"
                font.italic = True
                font.size = Pt(item["font"]["size"])
                font.color.rgb = RGBColor(21, 50, 158)

        docx_stream = BytesIO()
        self.document.save(docx_stream)
        docx_stream.seek(0)

        return docx_stream.getvalue()

